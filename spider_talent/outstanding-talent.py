import re
import os
import json
import math
import time
import requests
import pymysql
import docx
from win32com import client as wc
'''
1- 杭州市杰出人才数据爬取，从网站源数据到入库（包含了更新操作）
2- 符合要求的保存在本地，更新时比较
'''
headers = {
    "user-agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/83.0.4103.116 Safari/537.36",
}
proxys = []
def dl():
    dlurl = 'http://dynamic.goubanjia.com/dynamic/get/d490a5d4debefc8980ae6ee4c4148552.html?sep=3'
    resp = requests.get(dlurl).text
    time.sleep(2)
    resp = re.sub(r'\n', '', resp)
    proxy = {
        'http': resp
    }
    proxys.append(proxy)
    print(proxys[-1])
dl()

def get_path():
    rel_path = os.path.join('./', 'talent_files')
    return rel_path

def main():
    num = 0
    single = 1
    need_data = []
    data = {
        "startrecord": 1,
        "endrecord": 120,
        "perpage": 40,
        "col": 1,
        "appid": 1,
        "webid": 3163,
        "path": "/",
        "columnid": 1587845,
        "sourceContentType": 1,
        "unitid": 4840420,
        "webname": "杭州市人力资源和社会保障局",
        "permissiontype": 0,
    }
    url = "http://hrss.hangzhou.gov.cn/module/jpage/dataproxy.jsp?"
    response_web = ""
    for IP in range(10):
        try:
            response_web = requests.request(method='post', url=url, headers=headers, data=data, proxies=proxys[-1],
                                            timeout=15)
            # print(response)
            if response_web.status_code == 200:
                response_web = response_web.content.decode('utf8')
                break
            time.sleep(2)
        except Exception:
            dl()
            time.sleep(1.5)

    pattern = re.compile('<totalrecord>(.*?)</totalrecord>')
    flag = int(pattern.findall(response_web)[0])
    print('--------------第一次请求，总数--------------', flag)
    number = math.ceil(flag/120)

    for i in range(number):
        startrecord = num + 1
        endrecord = 120 * single

        if endrecord > flag:
            endrecord = flag
        data = {
            "startrecord": startrecord,
            "endrecord": endrecord,
            "perpage": 40,
            "col": 1,
            "appid": 1,
            "webid": 3163,
            "path": "/",
            "columnid": 1587845,
            "sourceContentType": 1,
            "unitid": 4840420,
            "webname": "杭州市人力资源和社会保障局",
            "permissiontype": 0,
        }
        url = "http://hrss.hangzhou.gov.cn/module/jpage/dataproxy.jsp?"

        response_web = ""
        for IP in range(10):
            try:
                response_web = requests.post(url=url, headers=headers, data=data, proxies=proxys[-1], timeout=15)
                if response_web.status_code == 200:
                    response_web = response_web.content.decode('utf8')
                    break
                time.sleep(2)
            except Exception:
                dl()
                time.sleep(1.5)
        num += 120
        single += 1

        pattern = re.compile('target="_blank" title="(.*?)">')
        results = pattern.findall(response_web)
        print(len(results))
        print((results))
        title = results

        pattern = re.compile('<a href="(.*?)" target="')
        results = pattern.findall(response_web)
        print(len(results))
        print((results))
        detail_url = results

        data = zip(title, detail_url)
        for j in data:
            # print(j)
            if "杰出创业人才培育计划" in j[0] and "拟入选对象公示的通知" in j[0]:
                need_data.append(j)
        time.sleep(1.5)
    test_detail(need_data)

def test_detail(need_data):
    # data = [
    #         ('关于杭州市大学生杰出创业人才培育计划第八批拟入选对象公示的通知', '/art/2020/7/7/art_1587845_50037544.html'),
    #         ('关于杭州市大学生杰出创业人才培育计划第七批拟入选对象公示的通知', '/art/2019/5/7/art_1587845_33987077.html'),
    #         ('关于杭州市大学生杰出创业人才培育计划第六批拟入选对象公示的通知', '/art/2018/5/30/art_1587845_28251879.html'),
    #         ('关于2016年杭州市杰出创业人才培育计划拟入选对象公示的通知', '/art/2017/1/6/art_1587845_28251722.html'),
    #         ('关于2013年杭州市杰出创业人才培育计划拟入选对象公示的通知', '/art/2013/12/10/art_1587845_28251214.html'),
    #         ('关于2012年杭州市杰出创业人才培育计划拟入选对象公示的通知', '/art/2012/11/23/art_1587845_28251158.html')
    #         ]
    urls = list(set([i[1] for i in need_data]))
    need = []

    for u in urls:
        url = 'http://hrss.hangzhou.gov.cn' + u
        print(url)
        for IP in range(100):
            try:
                response = requests.get(url=url, headers=headers, proxies=proxys[-1], timeout=15)
                if response.status_code == 200:
                    response = response.content.decode('utf8')
                    break
                time.sleep(2)
            except Exception:
                dl()
                time.sleep(1.5)
        # print(response)
        pattern = re.compile('附.*? href="(.*?)"', re.I)

        results = pattern.findall(response)[0]
        print(results)
        # print((results))
        need.append(results)
        time.sleep(1.5)
    print(need)
    # 判断本地是否有记录
    path = get_path()
    if not os.path.exists(path):
        parse_word(need, [])
    else:
        with open(os.path.join(path, 'record.json'), 'r')as f:
            content = json.loads(f.read())
            need = [i for i in need if i not in content]
            parse_word(need, content)

def parse_word(need, content):
    sum = []
    try:
        pa = get_path()
        if not os.path.exists(pa):
            os.mkdir(pa)
        new_need = need + content

        if len(need) == 0:
            print('------------------无数据更新-------------------')

        for n in need:
            url = "http://hrss.hangzhou.gov.cn" + n
            print(url)
            if "filename" in url:
                name = url.split('filename=')[-1]
            else:
                name = url.split('/')[-1]
            response = requests.get(url=url, headers=headers, proxies=proxys[-1]).content

            path = get_path()
            if not os.path.exists(path):
                os.mkdir(path)
            with open(os.path.join(path, "{}".format(name)), 'wb')as f:
                f.write(response)

            if name.split('.')[-1] == 'files':
                abspath = os.path.abspath('.').split('.')[0]
                full = os.path.join(abspath, 'talent_files')
                # win32接口：读文件 支持绝对路径；
                word = wc.Dispatch("Word.Application")
                doc = word.Documents.Open(os.path.join(full, "{}".format(name)))
                new_name = name.split('.')[0]+'.docx'
                doc.SaveAs(os.path.join(full, "{}".format(new_name)), 12)
                doc.Close()
                word.Quit()
                os.remove(os.path.join(full, "{}".format(name)))
                name = new_name

            data = docx.Document(os.path.join(path, "{}".format(name)))
            table = data.tables[0].rows[0].cells

            ta = data.tables

            s = [i.text for i in table]
            # 培育对象姓名   企业名称
            for index, i in enumerate(s):
                print(index, i)
                if '城区' in i:
                    chengqu = index
                if '对象' in i or '申报人' in i:
                    shenbaoduixiang = index
                if '名称' in i:
                    mingxheng = index

            for i in ta:
                rows = i.rows
                print(len(rows))
                for j in range(1, len(rows)):
                    dic = {}
                    name = ta[0].cell(j, shenbaoduixiang).text
                    company = ta[0].cell(j, mingxheng).text
                    area = ta[0].cell(j, chengqu).text
                    dic['name'] = name
                    dic['company'] = company
                    dic['area'] = area
                    dic['city'] = '杭州'
                    sum.append(dic)
            print(sum)

        time.sleep(2.3)
    except Exception as e:
        print(e)
    else:
        with open(os.path.join(pa, "record.json"), 'w')as f:
            f.write(json.dumps(new_need, ensure_ascii=False))
        write_db(sum)

def write_db(sum):

    db = pymysql.connect(host="localhost", user="root", password='123456', database="test", port=3306)
    cursor = db.cursor()
    content = sum
    for i in content:
        n = i.get('name')
        company = i.get('company')
        area = i.get('area')
        city = i.get('city')

        sql = """insert into outstandingtalent_copy1_copy1 (id,name,companyname,area,city)values(NULL,'{}','{}','{}','{}')""".format(
            n, company, area, city)
        cursor.execute(sql)
        db.commit()
    db.close()
    # exit()

if __name__ == '__main__':
    main()
