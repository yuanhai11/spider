import time

import docx

from sqlalchemy import Column, String, create_engine, Integer
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base

# 创建对象的基类:
Base = declarative_base()


# 定义User对象:
class Medicine(Base):
    # 表的名字:
    __tablename__ = 'spider_honor_company'

    # 表的结构:
    id = Column(Integer(), primary_key=True, autoincrement=True)
    company_name = Column(String(256))
    area = Column(String(256))
    license_code = Column(String(256))
    gmt_created = Column(String(256))
    gmt_updated = Column(String(256))
    company_id = Column(String(256))
    year = Column(String(256))
    types_of_honor = Column(String(256))
    batch = Column(String(256))


# 初始化数据库连接:
engine = create_engine('mysql+mysqlconnector://root:BOOT-xwork1024@192.168.2.97:3306/spider')
# 创建DBSession类型:
DBSession = sessionmaker(bind=engine)
# 创建session对象:
session = DBSession()


def main():
    sum = []
    data = docx.Document(
        r'D:\projects\S_Git_proj\spider\goverment_infomation_source_data\2022-09\2016-2022省科小名单\2017年度省级科技型中小企业备案名单.docx')
    print(data)
    all_tables = data.tables
    sum = []
    a = ''
    for ta in all_tables:
        rows = ta.rows
        for j in range(2, len(rows)):

            area = ta.cell(j, 1).text.strip()
            if area != '':
                area = area.split("（")[0]
                a = area
            company_name = ta.cell(j, 2).text
            license_code = ta.cell(j, 4).text

            year = '2017'
            # types_of_honor = '省科技型中小'
            # batch = '第三批'
            print(company_name, a, license_code)
            times = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())

            zhilian = Medicine(company_name=company_name,
                               license_code=license_code,
                               year=year, area=a,
                               # types_of_honor=types_of_honor,
                               # batch=batch,
                               gmt_updated=times, gmt_created=times,
                               )
            session.add(zhilian)
            session.commit()
            # sum.append(zhilian)
    # write_db(sum)



def main1():
    sum = []
    data = docx.Document(
        r'D:\projects\S_Git_proj\spider\goverment_infomation_source_data\2022-09\2016-2022省科小名单\2016年杭州市第三批浙江省科技型中小企业名单.docx')
    print(data)
    all_tables = data.tables
    sum = []
    print(all_tables)

    for index, ta in enumerate(all_tables):
        rows = ta.rows
        print(len(rows))

        for j in range(1, len(rows)):
            company_name = ta.cell(j, 1).text
            area = ta.cell(j, 2).text
            license_code = None
            year = '2016'
            types_of_honor = '省科技型中小'
            batch = '第三批'
            print(company_name, area, year, types_of_honor, batch)
            times = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())

            zhilian = Medicine(company_name=company_name,
                               license_code=license_code,
                               year=year, area=area,
                               types_of_honor=types_of_honor,
                               batch=batch,gmt_updated=times,gmt_created=times,
                               )
            sum.append(zhilian)
    write_db(sum)


def write_db(sum):
    for i in sum:
        session.add(i)
    session.commit()
    session.close()


def aa():
    # encoding=utf-8

    from win32com import client as wc
    w = wc.gencache.EnsureDispatch('word.Application')
    doc = w.Documents.Open(r"C:\Users\20945\Desktop\2016-2022省科小名单\2018年度省级科技型中小企业备案名单.doc")
    doc.SaveAs2(r"D:\projects\S_Git_proj\spider\goverment_infomation_source_data\2022-09\2016-2022省科小名单\2018年度省级科技型中小企业备案名单.docx", 12)



def pdf19():
    # 19年
    sum = []
    import pdfplumber
    import xlwt
    # path = "aaaaaa.PDF"  # 导入PDF路径
    pdf = pdfplumber.open(r"C:\Users\20945\Desktop\（高新处）关于公布2019年度浙江省科技型中小企业.pdf")

    area = ''
    for page in pdf.pages:
        if page == 0:
            continue
        # 获取当前页面的全部文本信息，包括表格中的文字
        # print(page.extract_text())
        for table in page.extract_tables():
            # print(table)
            for row in table:
                if row[0] == "" or row[4] == "":
                    continue
                try:
                    if row[1].replace("\n", "").strip() != '':
                        area = row[1].split("（")[0].replace("\n", "").strip()
                    company = row[2].replace("\n", "").strip()
                    code = row[3] + row[4]
                    year = '2019'
                    print(company, code, area)

                    times = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
                    zhilian = Medicine(company_name=company,
                                       license_code=code,
                                       year=year, area=area,
                                       gmt_updated=times, gmt_created=times,
                                       )
                    sum.append(zhilian)
                except Exception:
                    continue
    pdf.close()
    write_db(sum)

def pdf21():
    sum = []
    import pdfplumber
    import xlwt
    # path = "aaaaaa.PDF"  # 导入PDF路径
    pdf = pdfplumber.open(r"C:\Users\20945\Desktop\附件：2021年度省级科技型中小企业备案名单.pdf")

    area = ''
    flag = False
    for page in pdf.pages:
        if page == 0:
            continue
        # 获取当前页面的全部文本信息，包括表格中的文字
        # print(page.extract_text())
        for table in page.extract_tables():
            # print(table)
            for row in table:
                if row[0] == "序号":
                    continue
                if row[1]==None and flag!=True:
                    flag = True
                    continue
                if row[1] == None and flag:
                    area = row[0].split("(")[0].replace("\n", "").strip()
                    flag = False
                    continue

                try:
                    company = row[1].replace("\n", "").strip()
                    code = row[2]
                    year = '2021'
                    print(company, code, area)
                    times = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
                    zhilian = Medicine(company_name=company,
                                       license_code=code,
                                       year=year, area=area,
                                       gmt_updated=times, gmt_created=times,
                                       )
                    sum.append(zhilian)
                except Exception:
                    continue
    pdf.close()
    write_db(sum)

if __name__ == '__main__':
    # main1()
    # aa()
    # main()
    # pdf19()
    pdf21()