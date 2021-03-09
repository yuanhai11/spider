import xlrd
import pymysql
import docx
import os
# from win32com import client as wc


# a = os.listdir(r'C:\Users\20945\Desktop\Spider\now_spider\goverment_infomation_source_data\files')
# for i in a:
#     print(i)
# exit()

# def parse():
    # data = docx.Document(r'C:\Users\20945\Desktop\Spider\now_spider\goverment_infomation_source_data\files\杭州市2016年第一批浙江省科技型企业公示名单.docx')
    # print(data)
    # table = data.tables[0]
    # sum = []
    # rows = table.rows
    # print(len(rows))
    # # exit()
    # xixi = ''
    # flag = True
    # for j in range(1, len(rows)):
    #     area = table.cell(j,1).text
    #     name = table.cell(j,2).text
    #     year = '2008'
    #     type = '省科技型'
    #     batch = '第二批'
    #     # print((j,name,area,year,type))
    #     # sum.append((name,area))
    #     if area == name:
    #         xixi = '浙江'
    #         flag = False
    #         continue
    #     if flag:
    #         if area not in '\u3000':
    #             xixi = area.split('（')[0]
    #         else:
    #             area = xixi
    #
    #     area = xixi
    #     print((area,name))
    #     sum.append((name,area,year,type,batch))
    #     # exit()
    # print(sum)
    # exit()

# insert_db(sum)

# with open(r'C:\Users\20945\Desktop\a.csv',encoding='gbk')as fp:
#     content = fp.readlines()
# print(content)
# a = []
# for i in content[1:]:
#
#     data = i.split(',')
#     a.append((data[1],data[2],data[3],data[4]))
# insert_db(a)

# def doc_docx(path,i):
#     os.remove(os.path.join(path, "{}".format(i)))
#     word = wc.Dispatch("Word.Application")
#     doc = word.Documents.Open(os.path.join(path, "{}".format(i)))
#     new_name = i.split('.')[0]+'.docx'
#     doc.SaveAs(os.path.join(path, "{}".format(new_name)), 12)
#     doc.Close()
#     word.Quit()
#

#
# if __name__ == '__main__':
#     path = r'C:\Users\20945\Desktop\政策信息源-word'
#     data = os.listdir(path)
#     print(data)
#     for i in data:
#         if i.endswith('.doc'):
#             doc_docx(path,i)



# data = xlrd.open_workbook(r'C:\Users\20945\Desktop\Spider\now_spider\goverment_infomation_source_data\files\关于领取2017年新认定雏鹰青蓝企业证书、2017年省高成长科技型中小企业证书的通知.xls')
# print(data)
# table = data.sheets()[2]  # 通过索引顺序获取
# a = []
# for i in range(2,41):
#     a.append(table.cell(i, 1).value)
# print(a)
# print(len(a))
# insert_db(a)

import re
import time
import requests
import pymysql
from lxml import etree
from sqlalchemy import Column, String, create_engine,Integer
from sqlalchemy.orm import sessionmaker
from sqlalchemy.ext.declarative import declarative_base
# 创建对象的基类:
Base = declarative_base()
# 定义User对象:
class Medicine(Base):
    # 表的名字:
    __tablename__ = 'spider_company_makerspace'

    # 表的结构:
    id = Column(Integer(), primary_key=True,autoincrement=True)
    company_name = Column(String(256))
    space_name = Column(String(256))
    year = Column(String(256))
    area = Column(String(256))
    company_id = Column(String(256))
    types = Column(String(256))
# 初始化数据库连接:
engine = create_engine('mysql+mysqlconnector://root:BOOT-xwork1024@192.168.2.99:3306/spider')
# 创建DBSession类型:
DBSession = sessionmaker(bind=engine)
# 创建session对象:
session = DBSession()

def main():
    sum = []
    data = docx.Document(r'C:\Users\20945\Desktop\Spider\local_spider\goverment_infomation_source_data\files\关于杭州市2020年第二批浙江省科技型中小企业拟认定名单的公示.docx')
    print(data)
    # exit()
    all_tables = data.tables
    sum = []
    for ta in all_tables:
        rows = ta.rows
        for j in range(1, len(rows)):
            company_name = ta.cell(j, 1).text
            company_place = ta.cell(j, 2).text
            year = '2020'
            types_of_honor = '省科技型中小'
            batch = '第二批'
            from Other.post_es import get_company_id
            company_id = get_company_id(company_name)

            print(company_name,company_place,year,types_of_honor,batch,company_id)
            zhilian = Medicine(company_name=company_name, company_place=company_place, year=year, types_of_honor=types_of_honor, batch=batch,
                               company_id=company_id
                               )
            sum.append(zhilian)
    write_db(sum)
def main1():
    sum = []
    data = docx.Document(r'C:\Users\20945\Desktop\Spider\local_spider\goverment_infomation_source_data\files\关于2020年杭州市市级众创空间拟认定名单的公示.docx')
    print(data)
    # exit()
    all_tables = data.tables
    sum = []

    for index,ta in enumerate(all_tables):
        li = ['市级标准化众创空间','市级国际化示范众创空间','市级专业化示范众创空间']
        types = li[index]
        rows = ta.rows
        for j in range(1, len(rows)):
            space_name = ta.cell(j, 1).text
            company_name = ta.cell(j, 2).text
            area = ta.cell(j, 3).text
            year = '2020'
            types = types
            from Other.post_es import get_company_id
            company_id = get_company_id(company_name)

            print(space_name,company_name,area,year,types,company_id)
            zhilian = Medicine(company_name=company_name, space_name=space_name, year=year, area=area, types=types,
                               company_id=company_id
                               )
            sum.append(zhilian)
    write_db(sum)

def write_db(sum):
    for i in sum:
        session.add(i)
    session.commit()
    session.close()

if __name__ == '__main__':
    main1()